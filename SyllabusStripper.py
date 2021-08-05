import re
import PyPDF2
import textract
import pandas as pd

#from docx import Document
from dateutil.parser import parse
from datetime import datetime
from icalendar import Calendar, Event
from docx import Document

from os import listdir

from os.path import isfile, join

import tempfile
import os

import tabula


# Stores an array of regs
# Quite possibly the worst line of code I have ever written
# Is there a better way to do this, yes? Do I know how, no? Could I learn how? Yes. Do I care? No.
regArray = ["Jan[^0-9]+[0-3][0-9]", "Feb[^0-9]+[0-2][0-9]", "Mar[^0-9]+[0-3][0-9]", "Apr[^0-9]+[0-3][0-9]", "May[^0-9]+[0-3][0-9]", "Jun[^0-9]+[0-3][0-9]", "Jul[^0-9]+[0-3][0-9]", "Aug[^0-9]+[0-3][0-9]", "Sep[^0-9]+[0-3][0-9]", "Oct[^0-9]+[0-3][0-9]", "Nov[^0-9]+[0-3][0-9]", "Dec[^0-9]+[0-3][0-9]",
            "Jan[^0-9]+[0-9]", "Feb[^0-9]+[0-9]", "Mar[^0-9]+[0-9]", "Apr[^0-9]+[0-9]", "May[^0-9]+[0-9]", "Jun[^0-9]+[0-9]", "Jul[^0-9]+[0-9]", "Aug[^0-9]+[0-9]", "Sep[^0-9]+[0-9]", "Oct[^0-9]+[0-9]", "Nov[^0-9]+[0-9]", "Dec[^0-9]+[0-9]",
            "[0-1][0-9][/][0-3][0-9]", "[0-1][0-9][/][0-9]", "[0-9][/][0-3][0-9]", "[0-9][/][0-9]"]


# Gets a vague representation if string is a date
def is_date(string, fuzzy=False):
    """
    Return whether the string can be interpreted as a date.

    :param string: str, string to check for date
    :param fuzzy: bool, ignore unknown tokens in string if True
    """

    try:
        parse(string, fuzzy=fuzzy)
        return True

    except ValueError:
        return False


# Gets an Assignment contain a list of date and descriptions
def regexGetList(assignment, index):
    for u in range(len(regArray)):
        regDate = re.findall(regArray[u], str(assignment[index]))
        if (regDate != None):
            # Gets sublist from j to the end, ensures that the date will be the first indexreg
            for reg in regDate:
                # if u is 24, it is a date written in Month DD, so it changes that to MM/DD
                if(u < 24):
                    return list((parseDate(reg),)) + list(assignment[:index]) + list(assignment[index:])
                elif(is_date(reg)):  # Regular Output
                    return list((reg,)) + list(assignment[:index]) + list(assignment[index:])
    return None


# Parses a date from Month... Day format.
def parseDate(date):
    date = date.split(" ")
    if(len(date) == 1):
        date = [date[0], "5"]
    if(date[0].startswith("Jan")):
        return str("1/" + date[1])
    elif(date[0].startswith("Feb")):
        return str("2/" + date[1])
    elif(date[0].startswith("Mar")):
        return str("3/" + date[1])
    elif(date[0].startswith("Apr")):
        return str("4/" + date[1])
    elif(date[0].startswith("May")):
        return str("5/" + date[1])
    elif(date[0].startswith("Jun")):
        return str("6/" + date[1])
    elif(date[0].startswith("Jul")):
        return str("7/" + date[1])
    elif(date[0].startswith("Aug")):
        return str("8/" + date[1])
    elif(date[0].startswith("Sep")):
        return str("9/" + date[1])
    elif(date[0].startswith("Oct")):
        return str("10/" + date[1])
    elif(date[0].startswith("Nov")):
        return str("11/" + date[1])
    elif(date[0].startswith("Dec")):
        return str("12/" + date[1])


'''
    This methods manages files with a table, and parses info in a [Date,Description....] format
    It reads in values and uses a regex, to find valid dates
    Returns a list of Assignments with dates as the header

'''


def docxWithTable(document):
    # Stores rows and columns of a table
    data = []

    keys = None

    for table in document.tables:
        for i, row in enumerate(table.rows):
            text = (cell.text for cell in row.cells)

            # Establish the mapping based on the first row
            # headers; these will become the keys of our dictionary
            if i == 0:
                keys = tuple(text)
                continue

            # Construct a dictionary for this row, mapping
            # keys to values for this row
            row_data = tuple(text)
            data.append(row_data)

    datesArray = []

    for i in data:
        for j in range(len(i)):

            # checks if there is a valid instance of reg array in i[j]
            dateInfo = regexGetList(i, j)
            if(dateInfo != None):
                datesArray.append(dateInfo)
                break

    # Iterates through rows and columns, if it encounters a date, adds it to data array

    return datesArray


def getText(doc):
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)


def docxWithoutTable(document):
    text = getText(document)

    text = text.replace("\r", "")

    textArray = text.split("\n")

    for i in range(len(textArray)):
        textArray[i] = textArray[i].split(" ")
    datesArray = []

    for i in textArray:
        for j in range(len(i)):
            # checks if there is a valid instance of reg array in i[j]
            dateInfo = regexGetList(i, j)
            if(dateInfo != None):
                datesArray.append(dateInfo)
                break
    return datesArray


def docx(file):
    document = Document(file)

    if(len(document.tables) == 0):
        return docxWithoutTable(document)
    else:
        return docxWithTable(document)


'''
    This methods manages files with a table, and parses info in a [Date,Description....] format using tabula
    It reads in values and uses a regex, to find valid dates
    Returns a list of Assignments with dates as the header

'''


def pdfWithTable(dfOfInterest):
    o2dArray = pd.DataFrame.to_numpy(dfOfInterest)

    datesArray = []

    for i in o2dArray:
        for j in range(len(i)):

            # checks if there is a valid instance of reg array in i[j]
            dateInfo = regexGetList(i, j)
            if(dateInfo != None):
                datesArray.append(dateInfo)
                break

    # Iterates through rows and columns, if it encounters a date, adds it to data array

    # print(datesArray[0])

    return datesArray


def pdfWithoutTable(file):
    pdfFileObj = open(file, 'rb')

    pdfReader = PyPDF2.PdfFileReader(pdfFileObj)

    num_pages = pdfReader.numPages
    count = 0
    text = ""

    while count < num_pages:
        pageObj = pdfReader.getPage(count)
        count += 1
        text += pageObj.extractText()

    if text != "":
        text = text

    text = text.replace("\r", "")

    text = text.replace("\n", "")

    textArray = text.split(" ")

    dates = []

    for i in range(len(regArray)):
        shit = re.search(regArray[i], text)
        if shit != None:
            dates.append(shit)

    '''
    print(dates[0].span())
    print(text[dates[0].span()[0]:dates[0].span()[1]])
    
    print(dates[1].span())
    print(text[dates[1].span()[0]:dates[1].span()[1]])
    '''

    filteredDatesAndIndex = []

    for i in range(len(dates)):
        if dates[i].span()[1] - dates[i].span()[0] < 15:
            filteredDatesAndIndex.append(
                [text[dates[i].span()[0]:dates[i].span()[1]], dates[i].span()])

    for i in range(len(filteredDatesAndIndex)):
        # print(filteredDatesAndIndex[i][0])
        filteredDatesAndIndex[i][0] = parseDate(filteredDatesAndIndex[i][0])

    dates = []

    for i in range(len(filteredDatesAndIndex)):
        if filteredDatesAndIndex[i][0] != None:
            dates.append(filteredDatesAndIndex[i])

    descriptions = []

    for i in range(len(dates)):
        endOfDate = dates[i][1][1]
        descriptions.append(text[endOfDate+1:endOfDate + 100])

    '''
    currentYear = datetime.now().year
    
    for i in range(len(dates)):
        if("-" in dates[i][0]):
            dates[i][0] = dates[i][0].replace("-","/")
        dates[i][0] = str(currentYear) + "/" +  dates[i][0]
        dates[i][0] = parse(dates[i][0])
    '''

    datesArray = []

    for i in range(len(dates)):
        datesArray.append([dates[i][0], descriptions[i]])

    return datesArray


def pdf(file):

    # This boi's going to take all of the tables in the pdf to use later

    df = tabula.read_pdf(file, pages='all')

    # This abolsolute beast of a piece of code will take those sweet juicy
    # tables and find the largest one, we will assume the largest table is the
    # right one.

    largestSize = -1
    dfOfInterest = None

    for i in range(len(df)):
        currdf = df[i]
        if len(currdf) * len(currdf.columns) > largestSize:
            largestSize = len(currdf) * len(currdf.columns)
            dfOfInterest = currdf

    if(largestSize > 7):
        return pdfWithTable(dfOfInterest)
    else:
        return pdfWithoutTable(file)


def convertToICS(datesArray):
    dates = []
    descriptions = []

    # Assuming the description is one index over to the right of the date
    # (this is a valid assumption in something like a table) we can update
    # the dates and descriptions arrays appropriately

    for i in range(len(datesArray)):
        dates.append(datesArray[i][0])
        descriptions.append(datesArray[i][1:])

    # This cleans up the data a little bit, it removes the end from - onward
    # and replaces it with the current year. This should ensure it works for
    # any semester. It then runs it through dateutil's parse function to standardize
    # the layout of the data, it should be 'YYYY-MM-DD HH:MM:SS'

    currentYear = datetime.now().year

    for i in range(len(dates)):
        if("-" in dates[i]):
            dates[i] = dates[i].replace("-", "/")
        dates[i] = str(currentYear) + "/" + str(dates[i])
        dates[i] = parse(dates[i])

    # !!! Note: dates now contains the dates in format 'YYYY-MM-DD HH:MM:SS'

    # Initializes calendar class. Full documentation can be found here:
    # https://icalendar.readthedocs.io/en/latest/usage.html
    # !!! Note must include description and start time at least

    cal = Calendar()

    for i in range(len(dates)):
        event = Event()  # Starts a new event
        dateComponents = str(dates[i]).split("-")
        year = int(dateComponents[0])
        month = int(dateComponents[1])
        day = int(dateComponents[2][:2])
        # print(dates[i])
        # print(day)
        hour = 11
        minutes = 59
        seconds = 59
        # print(hour)

        # Formats date and time into something usable for event

        event.add('dtstart', datetime(
            year, month, day, hour, minutes, seconds))

        # Adds description
        descriptionTotal = ""

        for j in descriptions[i]:
            if j != None:
                descriptionTotal += str(j) + " "

        descriptionTotal.replace("nan", "")
        descriptionTotal.replace("\n", "").replace("\r", "")

        event.add('summary', descriptionTotal)

        # Can add more attributes to every event using afformentioned docs

        cal.add_component(event)  # Writes event to calendar

    f = open(os.path.join('download.ics'), 'wb')
    f.write(cal.to_ical())
    f.close()


def func():
    main()


def main():
    print("The main method is running")
    # Gets User Input From Website(Will be edited afterwards)
    # File name will be uploaded_file.pdf or .doc or .docx

    onlyfiles = [f for f in listdir("./") if isfile(join("./", f))]

    file = None

    print(onlyfiles)

    for i in range(len(onlyfiles)):
        if "uploaded_file" in onlyfiles[i]:
            file = onlyfiles[i]

    print(file)

    if file is None:
        return

    fileType = file[file.index("."):]

    datesArray = []

    if(fileType == ".pdf"):
        datesArray = pdf(file)
    elif(fileType == ".docx"):
        datesArray = docx(file)
    else:
        return

    convertToICS(datesArray)


if __name__ == "__main__":
    main()
